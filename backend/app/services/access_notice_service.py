"""Generation of access notices for guardians."""
from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
import secrets

from sqlalchemy.orm import Session

from app.core.security import hash_password, verify_password
from app.models import Aluno, Usuario
from app.services.accounts import ensure_responsavel_user, ensure_aluno_user


@dataclass(frozen=True)
class AccessNoticeResult:
    buffer: BytesIO
    filename: str
    row_count: int


class AccessNoticeService:
    SITE_URL = "https://gestao.colaboraedu.cloud/login/aluno"

    def __init__(self, session: Session):
        self.session = session

    def generate_class_docx(
        self,
        *,
        tenant_id: int,
        academic_year_id: int | None,
        turma: str,
        school_name: str,
        tipo: str = "responsavel",
    ) -> AccessNoticeResult:
        alunos = self._list_students(tenant_id, academic_year_id, turma)
        if not alunos:
            raise ValueError("Nenhum aluno ativo encontrado para a turma informada.")

        qr_buffer = None
        if tipo == "aluno":
            try:
                import qrcode

                qr = qrcode.QRCode(version=1, box_size=10, border=1)
                qr.add_data(self.SITE_URL)
                qr.make(fit=True)
                qr_img = qr.make_image(fill_color="black", back_color="white")
                qr_buffer = BytesIO()
                getattr(qr_img, "save")(qr_buffer, format="PNG")
                qr_buffer.seek(0)
            except ImportError:
                qr_buffer = None

        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt
        except ImportError as exc:  # pragma: no cover - environment/configuration guard
            raise RuntimeError("Dependência python-docx não instalada.") from exc

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        for index, aluno in enumerate(alunos):
            if index and index % 2 == 0:
                document.add_page_break()
            elif index % 2 == 1:
                separator = document.add_paragraph()
                separator.paragraph_format.space_before = Pt(4)
                separator.paragraph_format.space_after = Pt(4)
                run = separator.add_run("─" * 78)
                run.font.size = Pt(7)
            if tipo == "aluno":
                usuario, is_new_user = self._get_student_user(aluno)
            else:
                usuario, is_new_user = self._get_guardian_user(aluno)

            has_active_temporary_password = (
                usuario.must_change_password
                and usuario.is_active
                and usuario.deleted_at is None
                and not usuario.is_archived
            )
            if tipo == "aluno":
                password = aluno.matricula
            else:
                should_issue_password = is_new_user or not has_active_temporary_password
                password = self._temporary_password() if should_issue_password else None
            if password:
                usuario.password_hash = hash_password(password)
                usuario.must_change_password = True
                usuario.is_active = True
                usuario.deleted_at = None
                usuario.is_archived = False
            self.session.add(usuario)
            self._append_notice(
                document=document,
                school_name=school_name,
                aluno=aluno,
                username=usuario.username,
                password=password,
                align=WD_ALIGN_PARAGRAPH,
                font_size=Pt,
                tipo=tipo,
                qr_buffer=qr_buffer,
            )

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        safe_turma = "".join(ch if ch.isalnum() else "_" for ch in turma).strip("_") or "turma"
        prefix = "comunicados_aluno" if tipo == "aluno" else "comunicados_acesso"
        return AccessNoticeResult(
            buffer=buffer,
            filename=f"{prefix}_{safe_turma}.docx",
            row_count=len(alunos),
        )

    def _list_students(
        self,
        tenant_id: int,
        academic_year_id: int | None,
        turma: str,
    ) -> list[Aluno]:
        query = self.session.query(Aluno).filter(
            Aluno.tenant_id == tenant_id,
            Aluno.turma == turma,
            Aluno.is_archived.is_(False),
        )
        if academic_year_id:
            query = query.filter(Aluno.academic_year_id == academic_year_id)
        return query.order_by(Aluno.nome.asc()).all()

    def _get_guardian_user(self, aluno: Aluno) -> tuple[Usuario, bool]:
        usuario, initial_password = ensure_responsavel_user(self.session, aluno)
        if not usuario:
            raise RuntimeError(f"Não foi possível criar responsável para {aluno.nome}.")
        usuario.aluno_id = aluno.id
        usuario.matricula = aluno.matricula
        usuario.tenant_id = aluno.tenant_id
        usuario.role = "responsavel"
        return usuario, initial_password is not None

    def _get_student_user(self, aluno: Aluno) -> tuple[Usuario, bool]:
        usuario, initial_password = ensure_aluno_user(self.session, aluno)
        if not usuario:
            raise RuntimeError(f"Não foi possível criar usuário aluno para {aluno.nome}.")
        usuario.aluno_id = aluno.id
        usuario.matricula = aluno.matricula
        usuario.tenant_id = aluno.tenant_id
        usuario.role = "aluno"
        return usuario, initial_password is not None

    @staticmethod
    def _temporary_password(length: int = 8) -> str:
        alphabet = "abcdefghjkmnpqrstuvwxyz23456789"
        return "".join(secrets.choice(alphabet) for _ in range(length))

    def _append_notice(
        self,
        *,
        document,
        school_name: str,
        aluno: Aluno,
        username: str,
        password: str | None,
        align,
        font_size,
        tipo: str = "responsavel",
        qr_buffer: BytesIO | None = None,
    ) -> None:
        def paragraph(text: str = "", *, bold: bool = False, center: bool = False, size: int = 9):
            p = document.add_paragraph()
            p.alignment = align.CENTER if center else align.LEFT
            p.paragraph_format.space_after = font_size(1)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(text)
            run.bold = bold
            run.font.size = font_size(size)
            return p

        title = document.add_paragraph()
        title.alignment = align.CENTER
        title.paragraph_format.space_after = font_size(0)
        run = title.add_run(school_name)
        run.bold = True
        run.font.size = font_size(11)

        subtitle = document.add_paragraph()
        subtitle.alignment = align.CENTER
        subtitle.paragraph_format.space_after = font_size(3)
        
        if tipo == "aluno":
            subtitle_text = "Comunicado de Acesso do Aluno ao ColaboraEdu"
            welcome_text = "Querido(a) aluno(a),"
            body_text = (
                "O ColaboraEdu permite acompanhar boletins, notas, ocorrências, comunicados "
                "e informações escolares, aproximando você da sua escola."
            )
            access_title = "Dados de acesso do aluno:"
        else:
            subtitle_text = "Comunicado de Acesso ao ColaboraEdu"
            welcome_text = "Prezada família,"
            body_text = (
                "O ColaboraEdu permite acompanhar boletins, notas, ocorrências, comunicados "
                "e informações escolares, aproximando família e escola."
            )
            access_title = "Dados de acesso do responsável:"

        run = subtitle.add_run(subtitle_text)
        run.bold = True
        run.font.size = font_size(12)

        paragraph(welcome_text)
        paragraph(body_text)

        paragraph("Dados do aluno:", bold=True)
        paragraph(f"Nome: {aluno.nome}")
        paragraph(f"Matrícula: {aluno.matricula}    Turma: {aluno.turma}    Turno: {aluno.turno}")

        if tipo == "aluno" and qr_buffer:
            # 2-column table: credentials on left, QR Code on right
            table = document.add_table(rows=1, cols=2)
            table.autofit = False
            
            from docx.shared import Inches
            table.columns[0].width = Inches(4.5)
            table.columns[1].width = Inches(2.0)
            
            left_cell = table.cell(0, 0)
            left_cell.width = Inches(4.5)
            right_cell = table.cell(0, 1)
            right_cell.width = Inches(2.0)
            
            # Left cell content
            p0 = left_cell.paragraphs[0]
            p0.alignment = align.LEFT
            p0.paragraph_format.space_after = font_size(1)
            p0.paragraph_format.line_spacing = 1.0
            run = p0.add_run(access_title)
            run.bold = True
            run.font.size = font_size(9)
            
            def add_cell_paragraph(cell, text: str = "", *, bold: bool = False, size: int = 9):
                p = cell.add_paragraph()
                p.alignment = align.LEFT
                p.paragraph_format.space_after = font_size(1)
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run(text)
                run.bold = bold
                run.font.size = font_size(size)
                return p
            
            add_cell_paragraph(left_cell, f"Site: {self.SITE_URL}")
            add_cell_paragraph(left_cell, f"Escola: {school_name}")
            if password:
                add_cell_paragraph(left_cell, f"Usuário: {username}    Senha temporária: {password}", bold=True)
            else:
                add_cell_paragraph(left_cell, f"Usuário: {username}    Senha temporária: já emitida anteriormente", bold=True)
            
            # Right cell content (QR Code)
            right_p = right_cell.paragraphs[0]
            right_p.alignment = align.CENTER
            run = right_p.add_run()
            qr_buffer.seek(0)
            run.add_picture(qr_buffer, width=Inches(1.1))
        else:
            paragraph(access_title, bold=True)
            paragraph(f"Site: {self.SITE_URL}")
            paragraph(f"Escola: {school_name}")
            if password:
                paragraph(f"Usuário: {username}    Senha temporária: {password}", bold=True)
            else:
                paragraph(f"Usuário: {username}    Senha temporária: já emitida anteriormente", bold=True)

        warning = document.add_paragraph()
        if tipo == "aluno" and qr_buffer:
            warning.paragraph_format.space_before = font_size(4)
        warning.paragraph_format.space_after = font_size(2)
        warning.paragraph_format.line_spacing = 1.0
        warning.add_run("Importante: ").bold = True
        warning.add_run(
            "no primeiro acesso, será solicitada a troca da senha. Guarde estes dados com segurança."
            if password
            else "a senha temporária já foi emitida em um comunicado anterior e foi preservada."
        )
        for run in warning.runs:
            run.font.size = font_size(9)

        footer = document.add_paragraph()
        footer.alignment = align.CENTER
        footer.paragraph_format.space_after = font_size(1)
        footer.add_run("Atenciosamente,\nEquipe Escolar").bold = True
        for run in footer.runs:
            run.font.size = font_size(9)
